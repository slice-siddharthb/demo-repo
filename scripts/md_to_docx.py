"""Convert the services tables Markdown doc to a formatted .docx file."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_code_block(doc, code_text):
    """Add a shaded code block paragraph."""
    lines = code_text.strip().split("\n")
    for line in lines:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
        # light grey background on paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F0F0")
        pPr.append(shd)


def apply_inline_code(run_text, paragraph):
    """Split text on backtick spans and add runs with inline-code style."""
    parts = re.split(r"`([^`]+)`", run_text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run.font.size = Pt(10)


def apply_bold_and_code(text, paragraph):
    """Handle **bold**, `code`, and plain text in a single paragraph."""
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(10)


def parse_table(lines):
    """Parse a markdown table block into list of row lists."""
    rows = []
    for line in lines:
        if re.match(r"^\|[-| :]+\|$", line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def render_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            apply_bold_and_code(cell_text, p)
            if r_idx == 0:
                set_cell_bg(cell, "003366")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
                    run.font.size = Pt(9)
            else:
                bg = "F7F9FC" if r_idx % 2 == 0 else "FFFFFF"
                set_cell_bg(cell, bg)
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def build_doc(md_path, out_path):
    with open(md_path, "r") as f:
        content = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            i += 1

        # H2
        elif line.startswith("## ") and not line.startswith("### "):
            add_horizontal_rule(doc)
            p = doc.add_heading(line[3:].strip(), level=2)
            p.runs[0].font.color.rgb = RGBColor(0x00, 0x55, 0x99)
            i += 1

        # H3
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            p.runs[0].font.color.rgb = RGBColor(0x00, 0x66, 0xAA)
            i += 1

        # H4
        elif line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)
            i += 1

        # Code block
        elif line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1  # skip closing ```

        # Table
        elif line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            render_table(doc, rows)

        # Blockquote
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            apply_bold_and_code(line[2:].strip(), p)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1

        # Bullet list
        elif re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            apply_bold_and_code(line[2:].strip(), p)
            i += 1

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\. ", "", line).strip()
            apply_bold_and_code(text, p)
            i += 1

        # Horizontal rule
        elif line.strip() == "---":
            add_horizontal_rule(doc)
            i += 1

        # Empty line
        elif line.strip() == "":
            i += 1

        # Normal paragraph
        else:
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.space_after = Pt(6)
            apply_bold_and_code(line.strip(), p)
            i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_doc(
        "/workspace/docs/HOR-DSA_Services___druid_gold.services_orders___services_trackings.md",
        "/workspace/docs/HOR-DSA_Services___druid_gold.services_orders___services_trackings.docx",
    )
